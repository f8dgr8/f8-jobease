from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import anthropic
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re
import os

app = Flask(__name__)
CORS(app)

def create_docx(content, filename):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    lines = content.split('\n')
    
    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        if line.isupper() and len(line) > 3:
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        elif i == 0 or 'Olaluwoye' in line:
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(16)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('-') or line.startswith('•'):
            doc.add_paragraph(line[1:].strip(), style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route('/')
def home():
    return jsonify({"message": "F8 JobEase API is running"})

@app.route('/api/health')
def health():
    return jsonify({"status": "healthy"})

@app.route('/api/generate-resume', methods=['POST'])
def generate_resume():
    try:
        data = request.json
        client = anthropic.Anthropic(api_key=data.get('apiKey'))
        
        prompt = f"""Tailor this resume to this job.

BASE RESUME:
{data.get('baseResume')}

JOB DESCRIPTION:
{data.get('jobDescription')}

Create complete tailored resume. Extract keywords and rewrite all sections. Keep dates/companies exact.

At end add:
---METADATA---
JOB_TITLE: [job title]
COMPANY_NAME: [company name]

Output only resume."""

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}]
        )
        
        resume_text = message.content[0].text
        
        job_title = "Resume"
        company_name = "Company"
        
        if "---METADATA---" in resume_text:
            metadata = resume_text.split("---METADATA---")[1]
            title_match = re.search(r'JOB_TITLE:\s*(.+)', metadata)
            company_match = re.search(r'COMPANY_NAME:\s*(.+)', metadata)
            if title_match:
                job_title = title_match.group(1).strip()
            if company_match:
                company_name = company_match.group(1).strip()
            resume_text = resume_text.split("---METADATA---")[0].strip()
        
        filename = f"FO_{job_title.replace(' ', '_')}_{company_name.replace(' ', '_')}.docx"
        docx_buffer = create_docx(resume_text, filename)
        
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-cover-letter', methods=['POST'])
def generate_cover_letter():
    try:
        data = request.json
        client = anthropic.Anthropic(api_key=data.get('apiKey'))
        
        prompt = f"""Create professional cover letter.

RESUME: {data.get('baseResume')}
JOB: {data.get('jobDescription')}

Write 3-4 paragraph cover letter."""

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        docx_buffer = create_docx(message.content[0].text, "Cover_Letter")
        
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name="FO_Cover_Letter.docx"
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-questions', methods=['POST'])
def generate_questions():
    try:
        data = request.json
        client = anthropic.Anthropic(api_key=data.get('apiKey'))
        
        prompt = f"""Generate 20 interview questions for {data.get('jobRole')}.
5 behavioral, 5 technical, 5 situational, 5 general. Number 1-20."""

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        questions = []
        for i, line in enumerate(message.content[0].text.split('\n')):
            if line.strip():
                cat = "Behavioral" if i < 5 else "Technical" if i < 10 else "Situational" if i < 15 else "General"
                questions.append({
                    "id": i,
                    "question": line.strip().lstrip('0123456789. '),
                    "category": cat
                })
        
        return jsonify({"questions": questions})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/generate-study-guide', methods=['POST'])
def generate_study_guide():
    try:
        data = request.json
        client = anthropic.Anthropic(api_key=data.get('apiKey'))
        
        prompt = f"""Create interview study guide for {data.get('jobTitle')} at {data.get('company', 'company')}.

Include: role overview, technical skills, 15 questions, STAR examples, company research, topics to study, questions to ask, checklist."""

        message = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4000,
            messages=[{"role": "user", "content": prompt}]
        )
        
        filename = f"FO_Study_Guide_{data.get('jobTitle', 'Job').replace(' ', '_')}.docx"
        docx_buffer = create_docx(message.content[0].text, filename)
        
        return send_file(
            docx_buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=filename
        )
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)
